from django.shortcuts import get_object_or_404, redirect, render
from django.views.generic import ListView, DetailView, CreateView, UpdateView, View
from django.urls import reverse_lazy
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.template.loader import render_to_string
from django.contrib import messages
from django.utils import timezone
import json
from io import BytesIO
import zipfile
from datetime import datetime

# PDF generation imports
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
except ImportError:
    pass

# Word generation imports
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    pass

from core.models import FinalAppointmentWithDoctorModel, IllnessHistory, DiagnosisTemplate
from application.sanatorium.forms.final_app_form import FinalAppointmentWithDoctorForm


class FinalAppointmentEnhancedCreateOrUpdateView(LoginRequiredMixin, View):
    """Enhanced view that handles either creating a new final appointment or updating an existing one"""

    def get(self, request, *args, **kwargs):
        history_id = kwargs.get('history_id')
        illness_history = get_object_or_404(IllnessHistory, pk=history_id)

        # Check if a final appointment already exists
        try:
            existing_appointment = FinalAppointmentWithDoctorModel.objects.get(illness_history=illness_history)
            # If it exists, redirect to the detail view
            return redirect('final_appointment_enhanced_detail', pk=existing_appointment.pk)
        except FinalAppointmentWithDoctorModel.DoesNotExist:
            # If it doesn't exist, show the create view
            create_view = FinalAppointmentEnhancedCreateView.as_view()
            return create_view(request, *args, **kwargs)

    def post(self, request, *args, **kwargs):
        history_id = kwargs.get('history_id')
        illness_history = get_object_or_404(IllnessHistory, pk=history_id)

        # Check if a final appointment already exists
        try:
            existing_appointment = FinalAppointmentWithDoctorModel.objects.get(illness_history=illness_history)
            # If it exists, use the update view to handle the POST
            update_view = FinalAppointmentEnhancedUpdateView.as_view()
            return update_view(request, pk=existing_appointment.pk)
        except FinalAppointmentWithDoctorModel.DoesNotExist:
            # If it doesn't exist, use the create view to handle the POST
            create_view = FinalAppointmentEnhancedCreateView.as_view()
            return create_view(request, *args, **kwargs)


class FinalAppointmentEnhancedCreateView(LoginRequiredMixin, CreateView):
    model = FinalAppointmentWithDoctorModel
    form_class = FinalAppointmentWithDoctorForm
    template_name = 'sanatorium/doctors/final_appointment/enhanced_form.html'

    def get_success_url(self):
        return reverse_lazy('final_appointment_enhanced_detail', kwargs={'pk': self.object.pk})

    def form_valid(self, form):
        form.instance.illness_history_id = self.kwargs.get('history_id')
        form.instance.doctor = self.request.user
        form.instance.created_by = self.request.user
        form.instance.modified_by = self.request.user
        
        # Calculate BMI automatically
        if form.instance.height and form.instance.weight:
            height_m = form.instance.height / 100  # Convert cm to meters
            form.instance.imt = form.instance.weight / (height_m ** 2)
            
            # BMI interpretation
            if form.instance.imt < 18.5:
                form.instance.imt_interpretation = 1  # Underweight
            elif 18.5 <= form.instance.imt < 25:
                form.instance.imt_interpretation = 2  # Normal
            elif 25 <= form.instance.imt < 30:
                form.instance.imt_interpretation = 3  # Overweight
            else:
                form.instance.imt_interpretation = 4  # Obese
        
        messages.success(self.request, 'Заключительный прием успешно создан.')
        return super().form_valid(form)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        illness_history = get_object_or_404(IllnessHistory, pk=self.kwargs.get('history_id'))
        context['illness_history'] = illness_history
        context['history'] = illness_history
        context['patient'] = illness_history.patient
        context['active_page'] = {'final_appointment': 'active'}
        
        # Get all related appointments for summary
        context['consultations'] = {
            'cardiologist': illness_history.cardiologist_consulting.all(),
            'neurologist': illness_history.neurologist_consulting.all(),
            'repeated': illness_history.repeated_appointment.all(),
        }
        
        return context


class FinalAppointmentEnhancedUpdateView(LoginRequiredMixin, UpdateView):
    model = FinalAppointmentWithDoctorModel
    form_class = FinalAppointmentWithDoctorForm
    template_name = 'sanatorium/doctors/final_appointment/enhanced_form.html'

    def get_success_url(self):
        return reverse_lazy('final_appointment_enhanced_detail', kwargs={'pk': self.object.pk})

    def form_valid(self, form):
        form.instance.modified_by = self.request.user
        
        # Calculate BMI automatically
        if form.instance.height and form.instance.weight:
            height_m = form.instance.height / 100
            form.instance.imt = form.instance.weight / (height_m ** 2)
            
            # BMI interpretation
            if form.instance.imt < 18.5:
                form.instance.imt_interpretation = 1
            elif 18.5 <= form.instance.imt < 25:
                form.instance.imt_interpretation = 2
            elif 25 <= form.instance.imt < 30:
                form.instance.imt_interpretation = 3
            else:
                form.instance.imt_interpretation = 4
        
        messages.success(self.request, 'Заключительный прием успешно обновлен.')
        return super().form_valid(form)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['illness_history'] = self.object.illness_history
        context['history'] = self.object.illness_history
        context['patient'] = self.object.illness_history.patient
        context['active_page'] = {'final_appointment': 'active'}
        context['is_update'] = True
        
        # Get all related appointments for summary
        context['consultations'] = {
            'cardiologist': self.object.illness_history.cardiologist_consulting.all(),
            'neurologist': self.object.illness_history.neurologist_consulting.all(),
            'repeated': self.object.illness_history.repeated_appointment.all(),
        }
        
        return context


class FinalAppointmentEnhancedListView(LoginRequiredMixin, ListView):
    model = FinalAppointmentWithDoctorModel
    template_name = 'sanatorium/doctors/final_appointment/enhanced_list.html'
    context_object_name = 'appointments'
    paginate_by = 20

    def get_queryset(self):
        print("Fetching final appointments for list view")
        history_id = self.kwargs.get('history_id')
        return FinalAppointmentWithDoctorModel.objects.filter(illness_history_id=history_id)

    def get_context_data(self, **kwargs):
        
        context = super().get_context_data(**kwargs)
        context['illness'] = self.object.illness_history if hasattr(self, 'object') else None
        context['active_page'] = {'final_appointments': 'active'}
        context['search_query'] = self.request.GET.get('search', '')
        context['status_filter'] = self.request.GET.get('status', '')
        context['history'] = get_object_or_404(IllnessHistory, pk=self.kwargs.get('history_id'))

        return context


class FinalAppointmentEnhancedDetailView(LoginRequiredMixin, DetailView):
    model = FinalAppointmentWithDoctorModel
    template_name = 'sanatorium/doctors/final_appointment/enhanced_detail.html'
    context_object_name = 'appointment'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        appointment = self.object
        
        context['illness_history'] = appointment.illness_history
        context['history'] = appointment.illness_history
        context['patient'] = appointment.illness_history.patient
        context['active_page'] = {'final_appointment': 'active'}
        
        # Calculate BMI category
        if appointment.imt:
            if appointment.imt < 18.5:
                context['bmi_category'] = 'Недостаточная масса тела'
                context['bmi_color'] = 'info'
            elif 18.5 <= appointment.imt < 25:
                context['bmi_category'] = 'Нормальная масса тела'
                context['bmi_color'] = 'success'
            elif 25 <= appointment.imt < 30:
                context['bmi_category'] = 'Избыточная масса тела'
                context['bmi_color'] = 'warning'
            else:
                context['bmi_category'] = 'Ожирение'
                context['bmi_color'] = 'danger'
        
        # Get treatment history summary
        history = appointment.illness_history
        context['treatment_summary'] = {
            'consultations': {
                'cardiologist': history.cardiologist_consulting.all(),
                'neurologist': history.neurologist_consulting.all(),
                'repeated': history.repeated_appointment.all(),
            },
            'medications': getattr(history, 'medications', []),
            'procedures': getattr(history, 'procedures', []),
            'lab_tests': getattr(history, 'lab_tests', []),
        }
        
        # Calculate treatment duration
        context['treatment_duration'] = (timezone.now().date() - history.created_at.date()).days
        
        return context


@login_required
def export_final_appointment_pdf(request, pk):
    """Export final appointment to PDF format"""
    appointment = get_object_or_404(FinalAppointmentWithDoctorModel, pk=pk)
    
    # Create PDF buffer
    buffer = BytesIO()
    
    # Create PDF document
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    # Define styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30,
        alignment=TA_CENTER
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=12,
        spaceAfter=12,
        spaceBefore=20
    )
    
    normal_style = styles['Normal']
    
    # Build PDF content
    story = []
    
    # Header
    story.append(Paragraph("HAYAT MEDICAL CENTER", title_style))
    story.append(Paragraph("Заключительный прием лечащего врача", heading_style))
    story.append(Spacer(1, 12))
    
    # Patient information
    patient_data = [
        ['Пациент:', appointment.illness_history.patient.full_name],
        ['Дата рождения:', appointment.illness_history.patient.date_of_birth.strftime('%d.%m.%Y') if appointment.illness_history.patient.date_of_birth else 'Не указана'],
        ['Лечащий врач:', appointment.doctor.full_name if appointment.doctor else 'Не указан'],
        ['Дата приема:', appointment.created_at.strftime('%d.%m.%Y %H:%M')],
        ['Статус:', appointment.state],
    ]
    
    patient_table = Table(patient_data, colWidths=[2*inch, 4*inch])
    patient_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('BACKGROUND', (0, 0), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(patient_table)
    story.append(Spacer(1, 20))
    
    # Vital signs
    if any([appointment.height, appointment.weight, appointment.heart_beat, appointment.arterial_high, appointment.arterial_low]):
        story.append(Paragraph("Жизненные показатели", heading_style))
        
        vital_data = []
        if appointment.height:
            vital_data.append(['Рост:', f"{appointment.height} см"])
        if appointment.weight:
            vital_data.append(['Вес:', f"{appointment.weight} кг"])
        if appointment.imt:
            vital_data.append(['ИМТ:', f"{appointment.imt:.1f}"])
        if appointment.heart_beat:
            vital_data.append(['ЧСС:', f"{appointment.heart_beat} уд/мин"])
        if appointment.arterial_high and appointment.arterial_low:
            vital_data.append(['АД:', f"{appointment.arterial_high}/{appointment.arterial_low} мм рт.ст."])
        
        if vital_data:
            vital_table = Table(vital_data, colWidths=[2*inch, 4*inch])
            vital_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(vital_table)
            story.append(Spacer(1, 20))
    
    # Diagnosis
    if appointment.diagnosis.exists():
        story.append(Paragraph("Диагнозы", heading_style))
        for diagnosis in appointment.diagnosis.all():
            story.append(Paragraph(f"• {diagnosis.name}", normal_style))
        story.append(Spacer(1, 20))
    
    # Objective status
    if appointment.objective_status:
        story.append(Paragraph("Объективный статус", heading_style))
        story.append(Paragraph(appointment.objective_status, normal_style))
        story.append(Spacer(1, 20))
    
    # Treatment results
    story.append(Paragraph("Результаты лечения", heading_style))
    story.append(Paragraph(appointment.treatment_results, normal_style))
    story.append(Spacer(1, 20))
    
    # Summary
    if appointment.summary:
        story.append(Paragraph("Заключение врача", heading_style))
        story.append(Paragraph(appointment.summary, normal_style))
        story.append(Spacer(1, 20))
    
    # Footer
    story.append(Spacer(1, 50))
    story.append(Paragraph(f"Врач: _________________ {appointment.doctor.full_name if appointment.doctor else ''}", normal_style))
    story.append(Spacer(1, 20))
    story.append(Paragraph(f"Дата: {timezone.now().strftime('%d.%m.%Y')}", normal_style))
    
    # Build PDF
    doc.build(story)
    
    # Get PDF content
    pdf = buffer.getvalue()
    buffer.close()
    
    # Create response
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="final_appointment_{appointment.pk}_{appointment.illness_history.patient.full_name}.pdf"'
    response.write(pdf)
    
    return response


@login_required
def export_final_appointment_word(request, pk):
    """Export final appointment to Word format"""
    appointment = get_object_or_404(FinalAppointmentWithDoctorModel, pk=pk)
    
    # Create Word document
    doc = Document()
    
    # Add header
    header = doc.add_heading('HAYAT MEDICAL CENTER', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Заключительный прием лечащего врача', level=1)
    
    # Patient information
    doc.add_heading('Информация о пациенте', level=2)
    patient_table = doc.add_table(rows=5, cols=2)
    patient_table.style = 'Table Grid'
    
    patient_data = [
        ('Пациент:', appointment.illness_history.patient.full_name),
        ('Дата рождения:', appointment.illness_history.patient.date_of_birth.strftime('%d.%m.%Y') if appointment.illness_history.patient.date_of_birth else 'Не указана'),
        ('Лечащий врач:', appointment.doctor.full_name if appointment.doctor else 'Не указан'),
        ('Дата приема:', appointment.created_at.strftime('%d.%m.%Y %H:%M')),
        ('Статус:', appointment.state),
    ]
    
    for i, (label, value) in enumerate(patient_data):
        patient_table.cell(i, 0).text = label
        patient_table.cell(i, 1).text = value
    
    # Vital signs
    if any([appointment.height, appointment.weight, appointment.heart_beat, appointment.arterial_high, appointment.arterial_low]):
        doc.add_heading('Жизненные показатели', level=2)
        
        vital_data = []
        if appointment.height:
            vital_data.append(('Рост:', f"{appointment.height} см"))
        if appointment.weight:
            vital_data.append(('Вес:', f"{appointment.weight} кг"))
        if appointment.imt:
            vital_data.append(('ИМТ:', f"{appointment.imt:.1f}"))
        if appointment.heart_beat:
            vital_data.append(('ЧСС:', f"{appointment.heart_beat} уд/мин"))
        if appointment.arterial_high and appointment.arterial_low:
            vital_data.append(('АД:', f"{appointment.arterial_high}/{appointment.arterial_low} мм рт.ст."))
        
        if vital_data:
            vital_table = doc.add_table(rows=len(vital_data), cols=2)
            vital_table.style = 'Table Grid'
            
            for i, (label, value) in enumerate(vital_data):
                vital_table.cell(i, 0).text = label
                vital_table.cell(i, 1).text = value
    
    # Diagnosis
    if appointment.diagnosis.exists():
        doc.add_heading('Диагнозы', level=2)
        for diagnosis in appointment.diagnosis.all():
            doc.add_paragraph(f"• {diagnosis.name}")
    
    # Objective status
    if appointment.objective_status:
        doc.add_heading('Объективный статус', level=2)
        doc.add_paragraph(appointment.objective_status)
    
    # Treatment results
    doc.add_heading('Результаты лечения', level=2)
    doc.add_paragraph(appointment.treatment_results)
    
    # Summary
    if appointment.summary:
        doc.add_heading('Заключение врача', level=2)
        doc.add_paragraph(appointment.summary)
    
    # Footer
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph(f"Врач: _________________ {appointment.doctor.full_name if appointment.doctor else ''}")
    doc.add_paragraph(f"Дата: {timezone.now().strftime('%d.%m.%Y')}")
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Create response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="final_appointment_{appointment.pk}_{appointment.illness_history.patient.full_name}.docx"'
    response.write(buffer.getvalue())
    
    return response


@login_required
def get_patient_vitals_history(request, history_id):
    """Get patient's vitals history for charts"""
    history = get_object_or_404(IllnessHistory, pk=history_id)
    
    # Get all appointments with vitals
    appointments = []
    
    # Add initial appointment data if exists
    if hasattr(history, 'initial_appointment') and history.initial_appointment:
        initial = history.initial_appointment
        if any([initial.height, initial.weight, initial.heart_beat]):
            appointments.append({
                'date': initial.created_at.strftime('%d.%m.%Y'),
                'type': 'Первичный прием',
                'height': initial.height,
                'weight': initial.weight,
                'heart_beat': initial.heart_beat,
                'arterial_high': getattr(initial, 'arterial_high', None),
                'arterial_low': getattr(initial, 'arterial_low', None),
            })
    
    # Add repeated appointments
    for appointment in history.repeated_appointment.all():
        if any([getattr(appointment, 'height', None), getattr(appointment, 'weight', None), getattr(appointment, 'heart_beat', None)]):
            appointments.append({
                'date': appointment.created_at.strftime('%d.%m.%Y'),
                'type': 'Повторный прием',
                'height': getattr(appointment, 'height', None),
                'weight': getattr(appointment, 'weight', None),
                'heart_beat': getattr(appointment, 'heart_beat', None),
                'arterial_high': getattr(appointment, 'arterial_high', None),
                'arterial_low': getattr(appointment, 'arterial_low', None),
            })
    
    # Add final appointment if exists
    try:
        final = FinalAppointmentWithDoctorModel.objects.get(illness_history=history)
        appointments.append({
            'date': final.created_at.strftime('%d.%m.%Y'),
            'type': 'Заключительный прием',
            'height': final.height,
            'weight': final.weight,
            'heart_beat': final.heart_beat,
            'arterial_high': final.arterial_high,
            'arterial_low': final.arterial_low,
        })
    except FinalAppointmentWithDoctorModel.DoesNotExist:
        pass
    
    # Sort by date
    appointments.sort(key=lambda x: datetime.strptime(x['date'], '%d.%m.%Y'))
    
    return JsonResponse({'appointments': appointments})