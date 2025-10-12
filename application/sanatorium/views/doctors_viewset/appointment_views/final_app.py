from django.shortcuts import get_object_or_404, redirect
from django.views.generic import ListView, DetailView, CreateView, UpdateView, View
from django.urls import reverse_lazy
from django.contrib.auth.mixins import LoginRequiredMixin
from django.http import HttpResponse
from django.utils import timezone
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.models import FinalAppointmentWithDoctorModel, IllnessHistory
from application.sanatorium.forms.final_app_form import FinalAppointmentWithDoctorForm


class FinalAppointmentCreateOrUpdateView(LoginRequiredMixin, View):
    """View that handles either creating a new final appointment or updating an existing one"""

    def get(self, request, *args, **kwargs):
        history_id = kwargs.get('history_id')
        illness_history = get_object_or_404(IllnessHistory, pk=history_id)

        # Check if a final appointment already exists
        try:
            existing_appointment = FinalAppointmentWithDoctorModel.objects.get(illness_history=illness_history)
            # If it exists, redirect to the update view
            return redirect('final_appointment_update', pk=existing_appointment.pk)
        except FinalAppointmentWithDoctorModel.DoesNotExist:
            # If it doesn't exist, show the create view
            create_view = FinalAppointmentCreateView.as_view()
            return create_view(request, *args, **kwargs)

    def post(self, request, *args, **kwargs):
        history_id = kwargs.get('history_id')
        illness_history = get_object_or_404(IllnessHistory, pk=history_id)

        # Check if a final appointment already exists
        try:
            existing_appointment = FinalAppointmentWithDoctorModel.objects.get(illness_history=illness_history)
            # If it exists, use the update view to handle the POST
            update_view = FinalAppointmentUpdateView.as_view()
            return update_view(request, pk=existing_appointment.pk)
        except FinalAppointmentWithDoctorModel.DoesNotExist:
            # If it doesn't exist, use the create view to handle the POST
            create_view = FinalAppointmentCreateView.as_view()
            return create_view(request, *args, **kwargs)


class FinalAppointmentCreateView(LoginRequiredMixin, CreateView):
    model = FinalAppointmentWithDoctorModel
    form_class = FinalAppointmentWithDoctorForm
    template_name = 'sanatorium/doctors/appointments/final_app/form.html'

    def get_success_url(self):
        return reverse_lazy('final_appointment_detail',
                            kwargs={'pk': self.object.pk})

    def form_valid(self, form):
        form.instance.illness_history_id = self.kwargs.get('history_id')
        form.instance.doctor = self.request.user
        form.instance.created_by = self.request.user
        form.instance.modified_by = self.request.user
        return super().form_valid(form)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['illness_history'] = get_object_or_404(IllnessHistory, pk=self.kwargs.get('history_id'))
        context['history'] = get_object_or_404(IllnessHistory, pk=self.kwargs.get('history_id'))
        return context


class FinalAppointmentUpdateView(LoginRequiredMixin, UpdateView):
    model = FinalAppointmentWithDoctorModel
    form_class = FinalAppointmentWithDoctorForm
    template_name = 'sanatorium/doctors/appointments/final_app/form.html'

    def get_success_url(self):
        return reverse_lazy('final_appointment_detail',
                            kwargs={'pk': self.object.pk})

    def form_valid(self, form):
        form.instance.modified_by = self.request.user
        return super().form_valid(form)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['illness_history'] = self.object.illness_history
        context['history'] = self.object.illness_history
        return context


class FinalAppointmentListView(LoginRequiredMixin, ListView):
    model = FinalAppointmentWithDoctorModel
    template_name = 'sanatorium/doctors/appointments/final_app/list.html'
    context_object_name = 'appointments'

    def get_queryset(self):
        history_id = self.kwargs.get('history_id')
        return FinalAppointmentWithDoctorModel.objects.filter(illness_history_id=history_id)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['illness_history'] = get_object_or_404(IllnessHistory, pk=self.kwargs.get('history_id'))
        context['history'] = get_object_or_404(IllnessHistory, pk=self.kwargs.get('history_id'))
        context['history'] = get_object_or_404(IllnessHistory, pk=self.kwargs.get('history_id'))
        return context


class FinalAppointmentDetailView(LoginRequiredMixin, DetailView):
    model = FinalAppointmentWithDoctorModel
    template_name = 'sanatorium/doctors/appointments/final_app/detail.html'
    context_object_name = 'appointment'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['illness_history'] = self.object.illness_history
        context['history'] = self.object.illness_history
        return context


class FinalAppointmentDownloadView(LoginRequiredMixin, View):
    """View for downloading appointment details as a Word document"""

    def get(self, request, *args, **kwargs):
        appointment_id = kwargs.get('pk')
        appointment = get_object_or_404(FinalAppointmentWithDoctorModel, pk=appointment_id)

        # Create a new Word document
        doc = docx.Document()

        # Add a title
        title = doc.add_heading(f'Заключительный приём #{appointment.id}', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add date and time
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_paragraph.add_run(f'Дата: {timezone.now().strftime("%d.%m.%Y %H:%M")}')
        date_run.font.size = Pt(10)

        # Add patient information
        doc.add_heading('Информация о пациенте', level=2)
        patient_table = doc.add_table(rows=1, cols=2)
        patient_table.style = 'Table Grid'
        patient_table.columns[0].width = Inches(2)
        patient_table.columns[1].width = Inches(4)

        # Add patient data
        row = patient_table.rows[0].cells
        row[0].text = 'Пациент'
        row[1].text = appointment.illness_history.patient.full_name

        row = patient_table.add_row().cells
        row[0].text = 'Дата приёма'
        row[1].text = appointment.created_at.strftime("%d.%m.%Y %H:%M")

        row = patient_table.add_row().cells
        row[0].text = 'Врач'
        row[1].text = appointment.doctor.full_name

        row = patient_table.add_row().cells
        row[0].text = 'Cito'
        row[1].text = 'Нет'  # FinalAppointmentWithDoctorModel doesn't have cito field

        # Add physical indicators
        doc.add_heading('Физические показатели', level=2)
        indicators_table = doc.add_table(rows=1, cols=2)
        indicators_table.style = 'Table Grid'

        row = indicators_table.rows[0].cells
        row[0].text = 'Давление'
        row[1].text = f'{appointment.arterial_high}/{appointment.arterial_low}'

        row = indicators_table.add_row().cells
        row[0].text = 'ИМТ'
        row[1].text = str(appointment.imt) if appointment.imt else '--'

        row = indicators_table.add_row().cells
        row[0].text = 'Рост'
        row[1].text = f'{appointment.height} см' if appointment.height else '--'

        row = indicators_table.add_row().cells
        row[0].text = 'Вес'
        row[1].text = f'{appointment.weight} кг' if appointment.weight else '--'

        row = indicators_table.add_row().cells
        row[0].text = 'Пульс'
        row[1].text = f'{appointment.heart_beat} уд/мин' if appointment.heart_beat else '--'

        row = indicators_table.add_row().cells
        row[0].text = 'ИМТ интерпретация'
        row[1].text = str(appointment.imt_interpretation) if appointment.imt_interpretation else '--'

        # Add examination data
        doc.add_heading('Данные осмотра', level=2)

        doc.add_heading('Жалобы пациента', level=3)
        doc.add_paragraph(appointment.complaint if hasattr(appointment, 'complaint') and appointment.complaint else 'Не указано')

        doc.add_heading('Объективные данные', level=3)
        doc.add_paragraph(appointment.objective_data if hasattr(appointment, 'objective_data') and appointment.objective_data else 'Не указано')

        doc.add_heading('Объективный статус', level=3)
        doc.add_paragraph(appointment.objective_status if appointment.objective_status else 'Не указано')

        # Add conclusion
        doc.add_heading('Заключение', level=2)

        doc.add_heading('Диагноз', level=3)
        diagnoses = appointment.diagnosis.all()
        if diagnoses:
            for diag in diagnoses:
                doc.add_paragraph(diag.name, style='List Bullet')
        else:
            doc.add_paragraph('Не указано')

        doc.add_heading('Заключение', level=3)
        doc.add_paragraph(appointment.summary if appointment.summary else 'Не указано')

        doc.add_heading('Результаты лечения', level=3)
        doc.add_paragraph(appointment.treatment_results)

        # Prepare response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=appointment_{appointment.id}.docx'

        # Save document to response
        doc.save(response)

        return response
